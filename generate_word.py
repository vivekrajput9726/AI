from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Inches(8.27)   # A4
sec.page_height = Inches(11.69)
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)
sec.right_margin  = Inches(1)
sec.left_margin   = Inches(1.5)

# ── Default styles ───────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

def set_font(run, bold=False, size=12, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text='', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, italic=False, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.underline = underline
    return p

def add_heading(text, level=1, underline=True):
    sizes = {1: 14, 2: 13, 3: 12, 4: 12}
    p = add_para(text, bold=True, size=sizes.get(level, 12),
                 align=WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=12, space_after=6, underline=(underline and level <= 2))
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx+1]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or c_idx == len(row)-1 else WD_ALIGN_PARAGRAPH.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def page_break():
    doc.add_page_break()

def chapter_header(text):
    """Add right-aligned italic small header like sample doc"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.bold = True

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light grey background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E1E1E')
    pPr.append(shd)
    run.font.color.rgb = RGBColor(212, 212, 212)

# ════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════
add_para('A PROJECT REPORT', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40, space_after=6)
add_para('ON', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
add_para('SYNORA HEALTH', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
add_para('CODING HOUSE', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=12)
add_para('PROJECT REPORT SUBMITTED IN THE PARTIAL FULFILMENT FOR THE DEGREE', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('OF', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('MASTER OF SCIENCE', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('IN', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('ARTIFICIAL INTELLIGENCE', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('(MSc AI)', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('SUBMITTED', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('BY', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('RIYA SINGH (242024008) (2024-26)', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('SNEHA RAVAL (232024002) (2024-26)', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('UNDER THE GUIDANCE OF', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('DR. SUNIL KUMAR &', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('DR. PAPRI DAS', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
add_para('[Auro University Logo]', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=12)
add_para('AURO UNIVERSITY, SURAT, GUJARAT', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Academic Year – 2024-26', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  DECLARATION – RIYA SINGH
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_para('AURO UNIVERSITY', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
add_para('Integral & Transformational Learning', italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
add_heading('CERTIFICATE OF UNDERTAKING', 1)
add_heading('(DECLARATION)', 2)
add_para('I, Riya Singh, Enrollment No. 242024008, hereby declare that the project entitled "SYNORA HEALTH", undertaken at the School of Information Technology, is submitted in partial fulfillment of the requirements for the degree of Masters of Science in Artificial Intelligence.', space_before=8, space_after=8)
add_para('I further declare that this project is my original work and has not been submitted, either in part or in full, for the award of any other degree, diploma, associate ship, fellowship, or any other similar title in AURO University or any other University/Institution.', space_after=16)
add_para('Date: ____________', space_after=8)
add_para('Place: Surat, Gujarat', space_after=24)
add_para('Signature of the Student', bold=True, space_after=4)
add_para('Riya Singh', space_after=4)
add_para('Enrollment No.: 242024008', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  DECLARATION – SNEHA RAVAL
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_para('AURO UNIVERSITY', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
add_para('Integral & Transformational Learning', italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
add_heading('CERTIFICATE OF UNDERTAKING', 1)
add_heading('(DECLARATION)', 2)
add_para('I, Sneha Raval, Enrollment No. 232024002, hereby declare that the project entitled "SYNORA HEALTH", undertaken at the School of Information Technology, is submitted in partial fulfillment of the requirements for the degree of Masters of Science in Artificial Intelligence.', space_before=8, space_after=8)
add_para('I further declare that this project is my original work and has not been submitted, either in part or in full, for the award of any other degree, diploma, associate ship, fellowship, or any other similar title in AURO University or any other University/Institution.', space_after=16)
add_para('Date: ____________', space_after=8)
add_para('Place: Surat, Gujarat', space_after=24)
add_para('Signature of the Student', bold=True, space_after=4)
add_para('Sneha Raval', space_after=4)
add_para('Enrollment No.: 232024002', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  INSTITUTE CERTIFICATE
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_para('AURO UNIVERSITY', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
add_para('Integral & Transformational Learning', italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
add_heading('SCHOOL OF INFORMATION TECHNOLOGY', 1, underline=False)
add_heading('CERTIFICATE', 1, underline=False)
add_para('This is to certify that the project entitled "Synora Health" has been successfully carried out by the following students under our guidance in the partial fulfilment of the requirements for the degree of Masters of Science in Artificial Intelligence of AURO University, Surat, during the academic year 2024-26 (Semester IV).', space_before=8, space_after=12)
add_para('Students:', bold=True, space_after=4)
add_numbered('242024008 – Riya Singh')
add_numbered('232024002 – Sneha Raval')
add_para('Date: ____________', space_before=16, space_after=8)
add_para('Place: Surat', space_after=24)
add_para('Dr. Sunil Kumar & Dr. Papri Das', bold=True, space_after=4)
add_para('Project Guide / Supervisor', space_after=16)
add_para('Prof. Anurag Dixit', bold=True, space_after=4)
add_para('Head, School of Information Technology', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  COMPANY CERTIFICATE – RIYA
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_para('</> CODING HOUSE', bold=True, size=16, space_before=4, space_after=4)
add_para('+91 76009 78848 | support@codinghouse.in | www.codinghouse.in\n410 - Exito Commercial Hub, Surat, Gujarat | UDIYAM-GJ-22-0179228', size=9, space_after=8)
divider()
add_heading('CERTIFICATE', 1, underline=False)
add_para('TO WHOMSOEVER IT MAY CONCERN', bold=True, space_before=8, space_after=8)
add_para('This is to certify that Ms. Riya Singh, a student pursuing Master in Artificial Intelligence & Machine Learning, has successfully completed her training and project work at Coding House.', space_after=8)
add_para('During the period from 7th Jan 2026 to 1st May 2026, she actively worked on AI/ML-based projects including data analysis, model development, and implementation of machine learning algorithms. She worked on real-time AI applications and demonstrated practical knowledge of tools like Python, TensorFlow, and data modeling.', space_after=8)
add_para('She has also worked on Generative AI models and an AI-based Healthcare Platform (Synora Health), demonstrating strong analytical and problem-solving skills in real-world scenarios.', space_after=8)
add_para('She has worked under our guidance and direction. Her performance was found to be excellent. Throughout the training period, we found her to be hardworking, sincere, and highly dedicated.', space_after=8)
add_para('We wish her great success in all her future endeavors in the field of Artificial Intelligence and Machine Learning.', space_after=24)
add_para('Authorized Signatory', bold=True, space_after=20)
add_para('________________________', space_after=4)
add_para('Naynesh Patel', bold=True, space_after=4)
add_para('Founder & Tech Consultant, Coding House', space_after=0)
page_break()

# COMPANY CERTIFICATE – SNEHA
chapter_header('Synora Health')
add_para('</> CODING HOUSE', bold=True, size=16, space_before=4, space_after=4)
add_para('+91 76009 78848 | support@codinghouse.in | www.codinghouse.in\n410 - Exito Commercial Hub, Surat, Gujarat', size=9, space_after=8)
divider()
add_heading('CERTIFICATE', 1, underline=False)
add_para('TO WHOMSOEVER IT MAY CONCERN', bold=True, space_before=8, space_after=8)
add_para('This is to certify that Ms. Sneha Raval, a student pursuing Master in Artificial Intelligence & Machine Learning, has successfully completed her training and project work at Coding House.', space_after=8)
add_para('During the period from 7th Jan 2026 to 1st May 2026, she actively worked on AI/ML-based projects including data analysis, model development, and implementation of machine learning algorithms. She worked on real-time AI applications and demonstrated practical knowledge of tools like Python, TensorFlow, and data modeling.', space_after=8)
add_para('She has also worked on Generative AI models and an AI-based Healthcare Platform (Synora Health), demonstrating strong analytical and problem-solving skills in real-world scenarios.', space_after=8)
add_para('She has worked under our guidance and direction. Her performance was found to be excellent. Throughout the training period, we found her to be hardworking, sincere, and highly dedicated.', space_after=24)
add_para('Authorized Signatory', bold=True, space_after=20)
add_para('________________________', space_after=4)
add_para('Naynesh Patel', bold=True, space_after=4)
add_para('Founder & Tech Consultant, Coding House', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  NON-DISCLOSURE CERTIFICATE
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('NON-DISCLOSURE CERTIFICATE', 1)
divider()
add_para('This Non-Disclosure Certificate is issued in connection with the Major Project entitled "Synora Health" submitted by Riya Singh (242024008) and Sneha Raval (232024002), students of Master of Science in Artificial Intelligence at Auro University, Surat, for the academic year 2024-26.', space_before=8, space_after=8)
add_para('The project was carried out during the internship at Coding House, Surat, Gujarat under the supervision of Mr. Naynesh Patel, Founder & Tech Consultant.', space_after=12)
add_para('Terms of Non-Disclosure:', bold=True, space_after=4)
add_numbered('The proprietary source code, business logic, and internal architecture of the Synora Health platform shall not be disclosed to any third party without prior written consent from Coding House.')
add_numbered('The project documentation submitted for academic evaluation purposes is permitted under this agreement and shall be used solely for the purpose of degree completion at Auro University.')
add_numbered('Any publication, demonstration, or reproduction of the system shall require prior approval from both Coding House and Auro University.')
add_numbered('The students agree to provide a complete video demonstration of the system for viva and presentation evaluation, in accordance with university NDC guidelines.')
add_para('', space_before=16)
add_para('________________________         ________________________         ________________________', space_after=4)
add_para('Riya Singh                              Sneha Raval                           Naynesh Patel', bold=True, space_after=4)
add_para('Student (242024008)               Student (232024002)               Coding House', space_after=16)
add_para('Date: ____________', space_after=8)
add_para('Place: Surat, Gujarat', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENTS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('ACKNOWLEDGEMENT', 1)
add_para('I would like to express my sincere gratitude to Auro University for providing me with the opportunity to undertake the major project titled "Synora Health". This project has been a transformative learning experience that significantly enhanced my technical knowledge, practical skills, and understanding of real-world healthcare technology applications.', space_before=8, space_after=8)
add_para('I am deeply thankful to my internal project guide, Dr. Sunil Kumar, for his constant guidance, encouragement, and invaluable support throughout the development of this project. His expert suggestions, timely feedback, and constructive criticism played a significant role in the successful completion of this work.', space_after=8)
add_para('I would also like to extend my sincere gratitude to my external company mentor, Mr. Naynesh Patel, Founder & Tech Consultant at Coding House, Surat, for providing industry exposure, practical insights, and guidance during the internship period.', space_after=8)
add_para('I express my heartfelt thanks to the faculty members of the School of Information Technology for their continuous support and for providing the necessary resources and knowledge required for this project.', space_after=8)
add_para('I am also grateful to my colleague and project partner Sneha Raval for her collaboration, teamwork, and consistent dedication throughout the duration of this project.', space_after=20)
add_para('Riya Singh', bold=True, space_after=4)
add_para('Enrollment No.: 242024008', space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('ACKNOWLEDGEMENT', 1)
add_para('I express my sincere appreciation to Auro University for providing me with the opportunity to undertake the project "Synora Health". This project has been a wonderful learning experience that enhanced my understanding of healthcare technologies, artificial intelligence concepts, and real-world problem-solving skills.', space_before=8, space_after=8)
add_para('I am deeply thankful to my internal guide, Dr. Papri Das, for her valuable guidance, patience, and continuous support throughout the entire project. Her expert suggestions and encouragement inspired me to complete the work successfully and confidently.', space_after=8)
add_para('I would also like to convey my gratitude to Mr. Naynesh Patel of Coding House for sharing practical industry knowledge and mentoring us during the internship training period.', space_after=8)
add_para('My heartfelt thanks go to the faculty members of the School of Information Technology for their cooperation and for providing the academic support and resources necessary for completing this project.', space_after=8)
add_para('I am equally thankful to my project partner Riya Singh for her dedication, teamwork, and collaborative spirit that made this project a success.', space_after=20)
add_para('Sneha Raval', bold=True, space_after=4)
add_para('Enrollment No.: 232024002', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('TABLE OF CONTENTS', 1)
toc_data = [
    ('Sr.No', 'Chapters', 'Page No.'),
    ('1.', 'Introduction of the Project ("Synora Health")', '1'),
    ('', '1.1 Background of the project', '1'),
    ('', '1.2 Problem statement', '2'),
    ('', '1.3 Need and significance of the system', '3'),
    ('2.', 'Software/System Requirement Study', '4'),
    ('', '2.1 Hardware Requirements', '4'),
    ('', '2.2 Software Requirements', '5'),
    ('', '2.3 Tools and Technologies Used', '6'),
    ('3.', 'Objective of the Software', '9'),
    ('', '3.1 Primary and secondary objectives', '9'),
    ('', '3.2 Expected outcomes', '10'),
    ('4.', 'Feasibility Study', '11'),
    ('', '4.1 Technical Feasibility', '11'),
    ('', '4.2 Economic Feasibility', '13'),
    ('', '4.3 Operational Feasibility', '14'),
    ('5.', 'System Analysis', '15'),
    ('', '5.1 Functional Requirements', '15'),
    ('', '5.2 Non-functional Requirements', '17'),
    ('', '5.3 Existing system overview', '18'),
    ('6.', 'Software/System Design', '19'),
    ('', '6.1 Flowchart', '19'),
    ('', '6.2 Data Flow Diagram (Level 0 and Level 1)', '21'),
    ('', '6.3 Entity Relationship Diagram (ERD)', '23'),
    ('', '6.4 Use Case Diagram', '24'),
    ('', '6.5 Class Diagram', '26'),
    ('', '6.6 Sequence Diagram', '27'),
    ('7.', 'Front-End Screens', '29'),
    ('8.', 'Database Design', '49'),
    ('9.', 'Back-End Description', '55'),
    ('10.', 'Implementation', '61'),
    ('11.', 'Source Code', '69'),
    ('12.', 'Testing and Implementation', '77'),
    ('13.', 'Results and Discussion', '82'),
    ('14.', 'Limitations', '89'),
    ('15.', 'Future Enhancements', '91'),
    ('16.', 'Conclusion', '93'),
    ('17.', 'Bibliography / References', '98'),
    ('18.', 'Appendix', '101'),
]
tbl = add_table(['Sr.No', 'Chapters', 'Page No.'], toc_data[1:], [0.6, 4.5, 0.8])
page_break()

# ════════════════════════════════════════════════════════════
#  LIST OF FIGURES
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('LIST OF FIGURES', 1)
figs = [
    ('Figure 6.1.1','System Flowchart – Patient Side','19'),
    ('Figure 6.1.2','System Flowchart – Doctor Side','20'),
    ('Figure 6.2.1','DFD Level 0 – Context Diagram','21'),
    ('Figure 6.2.2','DFD Level 1 – Detailed Data Flow','22'),
    ('Figure 6.3.1','Entity Relationship Diagram','23'),
    ('Figure 6.4.1','Use Case Diagram – Patient Side','24'),
    ('Figure 6.4.2','Use Case Diagram – Doctor & Admin Side','25'),
    ('Figure 6.5.1','Class Diagram','26'),
    ('Figure 6.6.1','Sequence Diagram – Login Flow','27'),
    ('Figure 6.6.2','Sequence Diagram – Appointment Booking','28'),
    ('Figure 7.1.1','Synora Health Landing Page','29'),
    ('Figure 7.1.2','Login Page','30'),
    ('Figure 7.1.3','Registration Page with OTP Verification','31'),
    ('Figure 7.1.4','Patient Dashboard','32'),
    ('Figure 7.1.5','Doctor Listing Page','33'),
    ('Figure 7.1.6','Appointment Booking with Razorpay Payment','34'),
    ('Figure 7.1.7','AI Symptom Checker Module','35'),
    ('Figure 7.1.8','AI Health Copilot Chat Interface','36'),
    ('Figure 7.1.9','Health Records Management','37'),
    ('Figure 7.1.10','Medicine Reminder with Google Calendar','38'),
    ('Figure 7.1.11','Nearby Hospitals & Pharmacies','39'),
    ('Figure 7.1.12','Health Goals Tracker','40'),
    ('Figure 7.1.13','Insurance Policy & Claims Management','41'),
    ('Figure 7.1.14','Doctor Dashboard','42'),
    ('Figure 7.1.15','Doctor – Patient Management','43'),
    ('Figure 7.1.16','Digital Prescription Module','44'),
    ('Figure 7.1.17','Video Consultation – WebRTC Interface','45'),
    ('Figure 7.1.18','Admin Dashboard','46'),
    ('Figure 8.1.1','MongoDB Collections Overview','49'),
    ('Figure 11.1.1','FastAPI Application Entry Point – main.py','69'),
    ('Figure 11.1.2','Login Service – auth_service.py','70'),
    ('Figure 11.1.3','OTP Send Route – auth.py','71'),
    ('Figure 11.1.4','AI Symptom Analyzer – symptom_analyzer.py','72'),
    ('Figure 11.1.5','Payment Verification – payments.py','73'),
    ('Figure 11.1.6','WebSocket Chat Handler – chat.py','74'),
    ('Figure 11.1.7','Redux Auth Slice – authSlice.js','75'),
    ('Figure 11.1.8','Axios API Service with JWT Refresh – api.js','76'),
    ('Figure 18.1.1','Plagiarism Check Report (Paperpal + Turnitin)','101'),
]
add_table(['Figure No.','Figure Title','Page No.'], figs, [1.3, 4.1, 0.7])
page_break()

# LIST OF TABLES
chapter_header('Synora Health')
add_heading('LIST OF TABLES', 1)
tbls = [
    ('Table 2.1','Hardware Requirements','4'),
    ('Table 2.2','Software Requirements','5'),
    ('Table 2.3','Tools and Technologies Used','6'),
    ('Table 8.1','MongoDB Collections Overview','49'),
    ('Table 8.2','users Collection – Fields and Data Types','50'),
    ('Table 8.3','doctors Collection – Fields and Data Types','51'),
    ('Table 8.4','appointments Collection – Fields and Data Types','51'),
    ('Table 8.5','health_records Collection – Fields and Data Types','52'),
    ('Table 8.6','Primary and Foreign Key Relationships','53'),
    ('Table 8.7','Collection Relationships','54'),
    ('Table 10.1','Technologies Used in Synora Health','68'),
    ('Table 12.1','Unit Testing Results','77'),
    ('Table 12.2','Integration Testing Scenarios','79'),
    ('Table 12.3','System Testing Scenarios','80'),
    ('Table 12.4','User Acceptance Testing Observations','81'),
    ('Table 13.1','Appointment Booking System Output','83'),
    ('Table 13.2','AI Symptom Analysis Sample Output','84'),
    ('Table 13.3','Performance Discussion Summary','87'),
]
add_table(['Table No.','Table Title','Page No.'], tbls, [1.3, 4.1, 0.7])
page_break()

# LIST OF ABBREVIATIONS
chapter_header('Synora Health')
add_heading('LIST OF ABBREVIATIONS', 1)
abbr = [
    ('AI','Artificial Intelligence'),('API','Application Programming Interface'),
    ('CORS','Cross-Origin Resource Sharing'),('CRUD','Create, Read, Update, Delete'),
    ('CSS','Cascading Style Sheets'),('DFD','Data Flow Diagram'),
    ('EMR','Electronic Medical Record'),('ER','Entity Relationship'),
    ('GUI','Graphical User Interface'),('HTML','HyperText Markup Language'),
    ('HTTP','HyperText Transfer Protocol'),('HTTPS','HyperText Transfer Protocol Secure'),
    ('JWT','JSON Web Token'),('JS','JavaScript'),('JSON','JavaScript Object Notation'),
    ('LLM','Large Language Model'),('ML','Machine Learning'),
    ('MongoDB','Document-Oriented NoSQL Database'),('MVC','Model-View-Controller'),
    ('NDC','Non-Disclosure Compliance'),('NLP','Natural Language Processing'),
    ('NoSQL','Not Only SQL'),('OTP','One-Time Password'),('PDF','Portable Document Format'),
    ('REST','Representational State Transfer'),('SMS','Short Message Service'),
    ('SOS','Emergency Signal / Save Our Souls'),('UI','User Interface'),
    ('UX','User Experience'),('UUID','Universally Unique Identifier'),
    ('UAT','User Acceptance Testing'),('WebRTC','Web Real-Time Communication'),
    ('WS','WebSocket Protocol'),
]
add_table(['Abbreviation','Full Form'], abbr, [1.5, 4.7])
page_break()

# ════════════════════════════════════════════════════════════
#  CH 1: INTRODUCTION
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('1. INTRODUCTION OF THE PROJECT', 2)
add_heading('1.1 Background of the Project', 3)
add_para('The rapid advancement of digital health technologies and Artificial Intelligence has created an unprecedented opportunity to transform traditional healthcare delivery systems. In India, the healthcare sector faces significant challenges including insufficient doctor-patient ratios, fragmented medical record systems, inaccessible healthcare in rural areas, and lack of real-time health monitoring solutions.', space_before=8, space_after=8)
add_para('The "Synora Health" platform is developed to address these challenges by building a comprehensive AI-powered healthcare management ecosystem. The system integrates modern web technologies with advanced Artificial Intelligence and Large Language Models (LLMs) to provide a seamless experience for patients, doctors, and administrators in a single unified platform.', space_after=8)
add_para('The project was developed during a 6-month internship at Coding House, Surat, Gujarat, using a modern technology stack comprising Python FastAPI for the backend, React 18 with Vite for the frontend, MongoDB for data storage, and AI engines including Groq (Llama 3.1), Google Gemini, and OpenAI GPT-4 for intelligent health analysis.', space_after=8)
add_para('Synora Health serves three primary user roles — Patients, Doctors, and Administrators — each with a dedicated dashboard and feature set. The platform enables patients to book appointments, consult doctors via video calls, track health records, get AI-powered symptom analysis, manage medicine reminders, and monitor health goals.', space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('1.2 Problem Statement', 3)
add_para('In the current healthcare environment, both patients and medical professionals face numerous challenges due to the absence of an integrated, intelligent digital platform. The major problems identified are:', space_before=8, space_after=8)
add_bullet('Patients must visit multiple platforms separately for appointment booking, medical records, prescription management, and health monitoring.')
add_bullet('Manual appointment systems lead to scheduling conflicts, long waiting times, and poor patient experience.')
add_bullet('Medical records are stored in paper format or scattered across different systems, making it difficult to maintain a unified patient health history.')
add_bullet('Doctors lack AI-powered tools to assist in diagnosis, symptom analysis, and patient risk evaluation during consultations.')
add_bullet('Rural and semi-urban patients have limited access to specialist doctors due to geographic barriers, with no telemedicine solution available.')
add_bullet('Medicine adherence is poor due to the absence of intelligent reminder systems linked to digital calendars.')
add_bullet('Insurance claim management is manual, time-consuming, and lacks transparency for both patients and insurers.')
add_bullet('Emergency situations require instant access to medical assistance, which is not available through existing fragmented systems.')
add_bullet('Healthcare administrators lack real-time analytics and management tools to oversee platform activity and verify medical professionals.')
add_bullet('Existing platforms do not integrate AI-based health analysis, real-time consultation, payment gateways, and comprehensive health tracking into a single ecosystem.')
add_para('These limitations lead to delayed diagnoses, poor treatment adherence, inefficient healthcare delivery, and overall degraded patient outcomes.', space_before=8, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('1.3 Need and Significance of the System', 3)
add_para('The need for an intelligent, integrated healthcare platform like Synora Health is increasingly critical in today\'s fast-paced digital era.', space_before=8, space_after=8)
add_para('1. Unified Healthcare Ecosystem: Synora Health brings all healthcare services — appointments, records, consultations, AI analysis, payments, and monitoring — into a single platform.', space_after=6)
add_para('2. AI-Powered Health Intelligence: Integration of Large Language Models (Groq Llama 3.1, Google Gemini) enables real-time symptom analysis, health report interpretation, and personalized medical guidance with high accuracy.', space_after=6)
add_para('3. Telemedicine Accessibility: WebRTC-based video consultation makes specialist healthcare accessible regardless of geographic location.', space_after=6)
add_para('4. Secure Digital Health Records: A centralized, secure health records management system ensures patients and doctors always have access to accurate, up-to-date medical information.', space_after=6)
add_para('5. Real-Time Communication: WebSocket-based bidirectional chat enables continuous care coordination and timely medical guidance.', space_after=6)
add_para('6. Financial Integration: Razorpay payment gateway integration enables seamless appointment fee collection and insurance claim management.', space_after=6)
add_para('7. Preventive Healthcare: Features like medicine reminders, health goal tracking, vaccination management, and symptom diary promote proactive health management.', space_after=6)
add_para('8. Educational and Research Value: The project demonstrates practical implementation of AI, ML, real-time communication, and cloud technologies in a healthcare context.', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 2: SOFTWARE REQUIREMENTS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('2. SOFTWARE / SYSTEM REQUIREMENT STUDY', 2)
add_heading('2.1 Hardware Requirements', 3)
add_para('The hardware requirements specify the minimum and recommended system configuration needed for the successful development and execution of Synora Health.', space_before=8, space_after=8)
hw = [
    ('Processor','Intel Core i3 / AMD Ryzen 3','Intel Core i5/i7 or AMD Ryzen 5/7'),
    ('RAM','8 GB','16 GB or higher'),
    ('Storage','50 GB Free Space','256 GB SSD'),
    ('Operating System','Windows 10 / Ubuntu 20.04','Windows 11 / Ubuntu 22.04'),
    ('Internet Connection','10 Mbps','High-Speed 100 Mbps+'),
    ('Display Resolution','1366 × 768','Full HD (1920 × 1080)'),
    ('Webcam & Microphone','720p Camera','1080p HD Camera'),
    ('Browser','Chrome 90+ / Firefox 88+','Latest Chrome / Edge'),
]
add_table(['Component','Minimum Requirement','Recommended Requirement'], hw, [1.5, 2.4, 2.3])
add_para('Table 2.1: Hardware Requirements', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('2.2 Software Requirements', 3)
sw = [
    ('Python','3.11+','Core backend programming language'),
    ('FastAPI','0.109+','REST API framework for backend services'),
    ('Motor','3.3+','Async MongoDB driver for Python'),
    ('MongoDB','7.0','NoSQL document database'),
    ('Node.js','18+','JavaScript runtime for frontend tooling'),
    ('React','18.3','Frontend UI library'),
    ('Vite','5.4','Frontend build tool and dev server'),
    ('Tailwind CSS','3.4','Utility-first CSS framework'),
    ('Redux Toolkit','2.2','State management for React'),
    ('Groq SDK','Latest','Llama 3.1 AI model access'),
    ('Google Generative AI','Latest','Google Gemini 1.5 Pro access'),
    ('OpenAI SDK','1.56','GPT-4 AI access'),
    ('Twilio SDK','Latest','SMS OTP delivery service'),
    ('Razorpay SDK','1.4.1+','Payment gateway integration'),
    ('Uvicorn','Latest','ASGI server for FastAPI'),
    ('Pydantic','v2','Data validation and settings management'),
    ('PyJWT','Latest','JSON Web Token authentication'),
    ('bcrypt','Latest','Password hashing and verification'),
]
add_table(['Software/Tool','Version','Purpose'], sw, [1.7, 1.0, 3.5])
add_para('Table 2.2: Software Requirements', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('2.3 Tools and Technologies Used', 3)
tools = [
    ('IDE','Visual Studio Code','Primary development environment'),
    ('Version Control','Git + GitHub','Source code management'),
    ('API Testing','Postman / Swagger UI','REST API testing and documentation'),
    ('Database GUI','MongoDB Compass','MongoDB visual management'),
    ('Communication','WebSocket (FastAPI)','Real-time bidirectional chat'),
    ('Video Call','WebRTC','Peer-to-peer video consultation'),
    ('Maps','OpenStreetMap / Nominatim','Nearby hospitals and pharmacies'),
    ('Email Service','SMTP (Gmail)','OTP and notification emails'),
    ('SMS Service','Twilio','OTP delivery via SMS'),
    ('Payment','Razorpay','Appointment payment gateway'),
    ('Calendar','Google Calendar URL API','Medicine reminders integration'),
    ('Browser','Google Chrome / MS Edge','Testing and execution'),
    ('Package Manager','npm / pip','Dependency management'),
]
add_table(['Category','Tool / Technology','Purpose'], tools, [1.5, 1.8, 2.9])
add_para('Table 2.3: Tools and Technologies Used', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 3: OBJECTIVES
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('3. OBJECTIVE OF THE SOFTWARE', 2)
add_heading('3.1 Primary and Secondary Objectives', 3)
add_para('Primary Objectives:', bold=True, space_before=8, space_after=6)
add_para('The primary objective of "Synora Health" is to develop a fully integrated, AI-powered healthcare management platform that serves patients, doctors, and administrators through a unified digital ecosystem. The system aims to eliminate fragmentation in healthcare delivery by combining appointment management, medical records, real-time consultation, AI health analysis, payment processing, and health monitoring into a single platform.', space_after=8)
add_para('A key primary objective is to leverage Large Language Models (LLMs) including Groq Llama 3.1 and Google Gemini to provide intelligent symptom analysis, health report interpretation, and personalized medical guidance, making AI-enhanced healthcare accessible to all users.', space_after=8)
add_para('Secondary Objectives:', bold=True, space_after=6)
add_bullet('Implement secure OTP-based registration using Twilio SMS for authenticated access.')
add_bullet('Integrate Razorpay payment gateway for seamless appointment fee collection.')
add_bullet('Provide a comprehensive admin panel for platform monitoring, doctor verification, and analytics.')
add_bullet('Enable real-time chat between patients and doctors using WebSocket technology.')
add_bullet('Build a medicine reminder system integrated with Google Calendar.')
add_bullet('Develop a Smart Medical Report Analyzer using Google Gemini vision capabilities.')
add_bullet('Implement health tracking features: goals, vaccination records, family health, and symptom diary.')
add_bullet('Design responsive UI for seamless access on both desktop and mobile devices.')
add_bullet('Demonstrate practical application of AI/ML concepts in real-world healthcare scenarios.')
add_heading('3.2 Expected Outcomes', 3)
add_para('The expected outcome of "Synora Health" is a fully functional, production-ready AI-powered healthcare platform that significantly improves healthcare accessibility, efficiency, and patient outcomes. The system is expected to provide:', space_before=8, space_after=8)
add_bullet('A complete Patient Portal enabling appointment booking, health record management, AI symptom analysis, medicine reminders, and real-time doctor communication.')
add_bullet('A comprehensive Doctor Dashboard for patient management, appointment scheduling, prescription generation, and video consultations.')
add_bullet('An intelligent Admin Panel for platform governance, doctor verification, user management, and performance analytics.')
add_bullet('AI Health Copilot powered by Groq Llama 3.1 providing accurate medical guidance and health information.')
add_bullet('Smart Medical Report Analyzer that extracts and interprets health parameters from uploaded medical documents using Google Gemini.')
add_bullet('Secure payment processing for appointment bookings with support for multiple payment methods through Razorpay.')
page_break()

# ════════════════════════════════════════════════════════════
#  CH 4: FEASIBILITY
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('4. FEASIBILITY STUDY', 2)
add_para('The feasibility study of "Synora Health" was conducted to determine whether the proposed platform can be successfully developed, implemented, and maintained using available technologies, resources, and operational environments.', space_before=8, space_after=8)
add_heading('4.1 Technical Feasibility', 3)
add_para('Synora Health is technically feasible because it is built entirely on open-source, industry-standard technologies that are freely available, well-documented, and widely supported by active developer communities.', space_after=8)
add_para('The backend stack — Python 3.11, FastAPI, Motor, and MongoDB — represents the modern standard for building high-performance, scalable REST APIs with async capabilities. The frontend stack — React 18, Vite, and Tailwind CSS — enables rapid development of responsive, professional-grade user interfaces.', space_after=8)
add_para('Technical tools used in the project include:', space_after=6)
add_bullet('Python FastAPI for scalable, async REST API development')
add_bullet('MongoDB for flexible, schema-free document storage')
add_bullet('React 18 + Vite for high-performance frontend development')
add_bullet('Groq + Gemini APIs for AI health analysis')
add_bullet('WebRTC for browser-native video calling')
add_bullet('WebSocket for real-time bidirectional communication')
add_bullet('JWT for stateless, secure authentication')
add_bullet('Twilio SDK for SMS OTP delivery')
add_bullet('Razorpay SDK for payment processing')
add_heading('4.2 Economic Feasibility', 3)
add_para('Synora Health is economically feasible because the majority of tools, frameworks, libraries, and services used in its development are either free or available on generous free tiers. Python, FastAPI, React, Vite, Tailwind CSS, Redux Toolkit are all completely free and open-source. MongoDB offers a free 512MB Atlas cluster. Groq and Gemini APIs provide generous free-tier access. Twilio and Razorpay charge only for actual usage. The modular architecture ensures future upgrades can be done incrementally without significant additional cost.', space_before=8, space_after=0)
add_heading('4.3 Operational Feasibility', 3)
add_para('Synora Health is operationally feasible because it is designed with a user-first approach, featuring intuitive interfaces, role-based dashboards, and guided user flows that require no technical expertise to operate. The platform operates on standard web browsers without requiring any software installation by end users. The responsive design ensures functionality across desktop computers, tablets, and mobile phones.', space_before=8, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 5: SYSTEM ANALYSIS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('5. SYSTEM ANALYSIS', 2)
add_heading('5.1 Functional Requirements', 3)
add_para('Patient Functional Requirements:', bold=True, space_before=8, space_after=6)
patient_fr = ['User registration with SMS/Email OTP verification','Secure login with JWT-based session management',
    'Browse verified doctors by specialization and availability','Book appointments with Razorpay payment integration',
    'AI Symptom Checker powered by Groq Llama 3.1','AI Health Copilot for real-time medical guidance',
    'Smart Medical Report Analyzer using Google Gemini','Health records management (Lab Reports, Prescriptions, Diagnoses)',
    'Medicine reminders with Google Calendar integration','Health Goals tracking with progress monitoring',
    'Vaccination Tracker for immunization history','Family Health profiles management',
    'Symptom Diary for daily health logging','Insurance policy and claims management',
    'Nearby hospitals and pharmacies using OpenStreetMap','Emergency SOS alerts with location sharing',
    'Real-time chat with assigned doctors via WebSocket','Video consultation with doctors via WebRTC']
for fr in patient_fr: add_bullet(fr)
add_para('Doctor Functional Requirements:', bold=True, space_before=8, space_after=6)
doc_fr = ['Doctor registration with professional profile setup','Availability scheduling with time slot management',
    'View and manage pending/confirmed appointments','Access complete patient health history',
    'Write and manage digital prescriptions with AI assistance','Conduct video consultations via WebRTC',
    'Real-time chat with patients via WebSocket','View patient insurance details and claims',
    'Track earnings and payment history','Task and reminder management',
    'AI-assisted patient risk analysis']
for fr in doc_fr: add_bullet(fr)
add_para('Admin Functional Requirements:', bold=True, space_before=8, space_after=6)
adm_fr = ['Admin authentication and secure dashboard access','User management — view, activate, deactivate, delete',
    'Doctor verification and rejection workflow','Real-time platform statistics from MongoDB',
    'Appointment management and status control','Support ticket management',
    'Activity log monitoring and analytics reporting']
for fr in adm_fr: add_bullet(fr)
page_break()

chapter_header('Synora Health')
add_heading('5.2 Non-Functional Requirements', 3)
add_para('Performance: API response time under 500ms; AI analysis within 5 seconds; dashboard load within 2 seconds; video consultation stable at minimum 1 Mbps bandwidth.', space_before=8, space_after=8)
add_para('Security: All passwords bcrypt-hashed (cost factor 12); JWT access tokens expire in 60 minutes, refresh tokens in 7 days; CORS restricted to authorized origins; payment verification using HMAC-SHA256; role-based access control (RBAC) for all API endpoints; OTP expiry in 10 minutes.', space_after=8)
add_para('Reliability: System functions correctly when AI APIs are temporarily unavailable through fallback responses; MongoDB operations maintain ACID consistency; payment system prevents double-charging through idempotency verification.', space_after=8)
add_para('Scalability: FastAPI\'s async architecture supports high concurrent users; MongoDB\'s horizontal scaling supports large data volumes; modular frontend components enable independent feature scaling.', space_after=8)
add_para('Usability: Responsive design accessible on mobile, tablet, and desktop; intuitive navigation with clearly labeled sections; toast notifications for real-time user feedback; loading states and empty states for all data-driven components.', space_after=0)
add_heading('5.3 Existing System Overview', 3)
add_para('Existing healthcare digital platforms in India primarily offer isolated services without intelligent integration. Popular apps like Practo, Apollo 24/7, and mfine provide appointment booking, but they lack comprehensive AI-powered health analysis, real-time doctor-patient communication, and integrated health tracking.', space_before=8, space_after=8)
add_para('Key limitations of existing systems include: fragmentation requiring multiple apps, no real LLM integration (only rule-based chatbots), limited telemedicine (premium-only), no insurance integration, no preventive health tools in integrated form, poor doctor experience tools, and no smart report analysis. Synora Health overcomes all these limitations by delivering a fully integrated, AI-powered healthcare ecosystem.', space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 6: SYSTEM DESIGN
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('6. SOFTWARE / SYSTEM DESIGN', 2)
add_heading('6.1 Flowchart', 3)
add_para('[Figure 6.1.1 – System Flowchart Patient Side]\nSTART → User Registered? → No: Register Account → OTP Verification → Login | Yes: Login → Patient Dashboard → Select Feature (Book Appointment / AI Symptom / Health Records / Video Consult / Medicine Reminder / Insurance) → Logout/Continue → END', space_before=8, space_after=8, italic=True)
add_para('[Figure 6.1.2 – System Flowchart Doctor Side]\nSTART → Doctor Login (JWT) → Doctor Dashboard → Select Feature (Appointments / Patients / Prescriptions / Video Call / Messages) → Schedule/Availability → Earnings & Analytics → END', space_after=0, italic=True)
add_para('Insert actual flowchart diagrams here. Diagrams are available in the HTML report at c:\\Users\\Dell\\AI\\SYNORA_REPORT.html', bold=True, space_before=6, space_after=0)
add_heading('6.2 DFD Level 0 & Level 1', 3)
add_para('[Figure 6.2.1 – DFD Level 0 Context Diagram]\nExternal Entities: Patient User, Doctor User, Admin User ↔ Synora Health Platform (Central System). Data flows: Patient sends Login Data, Financial Input, Queries; receives Results & AI Responses. Doctor sends/receives Patient Info, Prescriptions. Admin sends Admin Commands; receives Reports & User Data.', space_before=8, space_after=8, italic=True)
add_para('[Figure 6.2.2 – DFD Level 1 Detailed Flow]\nProcesses: 1.0 Auth Module, 2.0 Appointment, 3.0 AI Engine, 4.0 Health Records, 5.0 Doctor Management, 6.0 Admin Panel, 7.0 Payment. Datastore: MongoDB DB connecting all processes.', space_after=0, italic=True)
add_heading('6.3 Entity Relationship Diagram (ERD)', 3)
add_para('[Figure 6.3.1 – Entity Relationship Diagram]\nMain Entities: USER (PK: _id, full_name, email, role, is_verified) → DOCTOR (PK: _id, FK: user_id, specialization, availability[], consultation_fee) [1:1]; USER → APPOINTMENT (PK: _id, FK: patient_id, doctor_id, status, payment_id) [1:M]; USER → HEALTH_RECORD (PK: _id, FK: patient_id, record_type, description) [1:M]; USER → MESSAGE (PK: _id, FK: sender_id, room_id, content) [1:M]; USER → PAYMENT (PK: _id, FK: user_id, razorpay_id, amount) [1:M]; USER → INSURANCE (PK: _id, FK: patient_id, policy_name, claim_status) [1:M]; APPOINTMENT → PAYMENT [1:1]', space_before=8, space_after=0, italic=True)
page_break()

chapter_header('Synora Health')
add_heading('6.4 Use Case Diagram', 3)
add_para('[Figure 6.4.1 – Use Case Diagram Patient Side]\nActor: Patient. Use Cases within Synora Health Patient Portal boundary: Register/OTP Verify, Login/Logout, Book Appointment, Make Payment (Razorpay), AI Symptom Checker, AI Health Copilot, Smart Report Analyzer, View Health Records, Medicine Reminder, Video Consultation, Chat with Doctor, Health Goals/Tracker, Nearby Hospitals, Emergency SOS.', space_before=8, space_after=8, italic=True)
add_para('[Figure 6.4.2 – Use Case Diagram Doctor & Admin Side]\nDoctor Actor: Dashboard, View Appointments, Manage Patients, Write Prescriptions, Video Call, Chat, Manage Schedule. Admin Actor: Login, Verify Doctors, Manage Users, View Analytics, Monitor Activity, Support Tickets.', space_after=0, italic=True)
add_heading('6.5 Class Diagram', 3)
add_para('[Figure 6.5.1 – Class Diagram]\nUser (base class: _id, full_name, email, role, is_verified + login(), logout()); Patient extends User (health_records, appointments, medicines + bookAppointment(), viewRecords()); Doctor extends User (specialization, availability, consultation_fee + writePrescription(), manageSchedule()); Admin extends User (admin_level, permissions + manageUsers(), verifyDoctor(), viewAnalytics()); Appointment (patient_id:ref, doctor_id:ref, status + updateStatus()); AIEngine (model, api_key + analyzeSymptoms(), analyzeReport())', space_before=8, space_after=0, italic=True)
add_heading('6.6 Sequence Diagram', 3)
add_para('[Figure 6.6.2 – Appointment Booking Sequence Diagram]\nLifelines: Patient UI → FastAPI → MongoDB → Razorpay → Doctor UI\n1. Patient selects doctor + slot → 2. FastAPI checks availability in MongoDB → 3. Slots confirmed → 4. Patient books → 5. FastAPI creates Razorpay order → 6. order_id returned → 7. Razorpay checkout shown → 8. [User Pays] → 9. HMAC verification → 10. Appointment saved to MongoDB → 11. Booking confirmed to patient → 12. Doctor notified via WebSocket', space_before=8, space_after=0, italic=True)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 7: FRONT-END SCREENS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('7. FRONT-END SCREENS', 2)
add_para('The front-end of "Synora Health" is designed to provide a modern, intuitive, and fully responsive interface for Patients, Doctors, and Administrators. The UI is built using React 18, Vite, Tailwind CSS, and Redux Toolkit. All screens feature gradient designs, real-time feedback via toast notifications, and role-based navigation through a persistent sidebar.', space_before=8, space_after=8)
add_heading('7.1 User Interface Screens', 3)

screens = [
    ('Landing Page (Figure 7.1.1)', 'Shows Synora Health hero section with "AI-Powered Healthcare" tagline, Get Started button, feature highlights (AI Diagnosis, Video Consult, Smart Records), and navigation bar with Login/Register links.'),
    ('Login Screen (Figure 7.1.2)', 'Email and password fields, Show/Hide password toggle, Sign In button, Quick Demo Login buttons (Patient / Doctor / Admin), "Don\'t have an account? Create one" link, Synora Health branding.'),
    ('Registration & OTP (Figure 7.1.3)', 'Step 1: Full Name, Email, Phone (+91 prefix), Password, Role selector. Step 2: 6-digit OTP input boxes, OTP channel indicator (SMS/Email), Demo OTP display, Resend OTP button.'),
    ('Patient Dashboard (Figure 7.1.4)', 'Good Morning/Afternoon/Evening greeting, Today\'s date, Upcoming appointment card, Quick action tiles, Health stats row, Health Goals card, Family Health card, Health Timeline card with gradient headers.'),
    ('Doctor Listing (Figure 7.1.5)', 'Searchable doctor cards with photo, name, specialization, experience, rating stars, consultation fee, availability day chips, Verified badge, "Book Now" button.'),
    ('Appointment Booking (Figure 7.1.6)', 'Calendar date picker, Time slot chip grid, Appointment type selector, Symptoms text area, "Pay & Book Appointment" button, Razorpay payment modal with UPI/Card/NetBanking tabs.'),
    ('AI Symptom Checker (Figure 7.1.7)', 'Symptom input, Body area selector, Duration, Severity slider. Results: conditions with probability, urgency level, recommended actions, "Take Photo of Report" camera button.'),
    ('AI Health Copilot (Figure 7.1.8)', 'Chat interface with Synora AI, message bubbles, Quick prompt chips, role-aware (Blue for Patient, Violet for Doctor), message history.'),
    ('Health Records (Figure 7.1.9)', 'Records list with type badges, search/filter, Add Record modal with type-specific forms (Lab Report, Prescription, Diagnosis, Other).'),
    ('Medicine Reminder (Figure 7.1.10)', 'Medicine cards with name, dosage, frequency, CalendarPlus icon per medicine, "Add All to Calendar" button, Active/Inactive toggle.'),
    ('Nearby Hospitals (Figure 7.1.11)', 'OpenStreetMap with pins, category filter tabs, distance selector, results list with name, address, distance, "Get Directions".'),
    ('Health Goals (Figure 7.1.12)', 'Multi-step wizard (5 steps), required Target/Current/Unit fields, progress bars, clickable step navigation bar, achievement badges.'),
    ('Insurance (Figure 7.1.13)', 'Policies list, Claims with filter pills (All/Pending/Approved/Rejected), claim cards with status badge, status update dropdown.'),
    ('Doctor Dashboard (Figure 7.1.14)', 'Teal gradient greeting banner with doctor name, appointment stats, Today\'s Schedule, Quick Actions, Upcoming & Recent Appointments.'),
    ('Patient Management (Figure 7.1.15)', 'Patient list, Patient Full Modal with tabs (Overview, Insights, Records, Insurance), AI health insights, consultation buttons.'),
    ('Prescription Module (Figure 7.1.16)', 'Patient selector, medication table, AI suggestions panel, digital prescription preview, print/download button.'),
    ('Video Consultation (Figure 7.1.17)', 'WebRTC video with local/remote streams, audio/video controls, screen share, in-call chat panel, call duration timer.'),
    ('Admin Dashboard (Figure 7.1.18)', 'Time-based greeting, real-time stats cards (clickable), Platform Growth chart, User Distribution donut chart, Platform Overview card, Doctor Verifications section.'),
]
for title, desc in screens:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(desc)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    add_para('[Insert Screenshot Here]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
page_break()

chapter_header('Synora Health')
add_heading('7.2 Forms and Navigation', 3)
add_para('Synora Health follows a sidebar-based navigation structure with role-specific menus. After successful authentication, users are redirected to their role-specific dashboard. The navigation system uses React Router v6 with protected routes that enforce role-based access control.', space_before=8, space_after=8)
add_para('Patient Navigation: Dashboard, Appointments (Doctor Listing), Health Records, My Records, Reports & Analysis, Medicine Reminder, Health Goals, Family Health, Health Timeline, Insurance, Nearby Hospitals, AI Health Copilot, Emergency SOS.', space_after=8)
add_para('Doctor Navigation: Dashboard, Appointments (with pending count badge), Patients, Consultations, Prescriptions, Medical Records, Reports & Analytics, Schedule, Messages (with unread badge), Earnings, Payments, Notifications, Referrals, Task & Reminders.', space_after=8)
add_para('Admin Navigation: Dashboard, Users, Doctors, Appointments, AI Analytics, Finance, Emergency Alerts, System Management, Notifications, Activity Log, Support Tickets, Reports, Settings.', space_after=0)
add_heading('7.3 Explanation of Each Screen', 3)
explanations = [
    ('7.3.1 Landing Page', 'Introduces Synora Health with professional hero section, key feature highlights, and clear call-to-action buttons for registration and login.'),
    ('7.3.2 Login Page', 'Provides secure login with JWT token management, role-based redirection, and Quick Demo Login buttons for testing all three roles.'),
    ('7.3.3 Registration Page', 'Two-step registration with OTP verification via Twilio SMS (email fallback), auto-login after successful verification.'),
    ('7.3.4 Patient Dashboard', 'Central hub with time-based greeting, health overview stats, upcoming appointment card, and quick access to all platform features.'),
    ('7.3.5 Appointment Booking', 'Complete booking flow with date picker, slot chips, Razorpay payment, and WebSocket notification to doctor on confirmation.'),
    ('7.3.6 AI Symptom Checker', 'Groq Llama 3.1 powered analysis providing conditions, urgency levels, action recommendations, and optional Gemini image analysis.'),
    ('7.3.7 AI Health Copilot', 'Role-aware LLM chat with different prompts for doctor (clinical) vs patient (health guidance), quick-action chips.'),
    ('7.3.8 Health Records', 'Type-specific forms (Lab Report, Prescription, Diagnosis) with search/filter and Gemini AI analysis integration.'),
    ('7.3.9 Video Consultation', 'WebRTC peer-to-peer video with audio/video controls, screen sharing, and in-call text chat panel.'),
    ('7.3.10 Admin Dashboard', 'Real-time stats from /admin/stats API, clickable stat cards navigating to relevant tabs, doctor verification workflow.'),
]
for title, desc in explanations:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': '); r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(desc); r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 8: DATABASE DESIGN
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('8. DATABASE DESIGN', 2)
add_para('Synora Health uses MongoDB 7.0 as its primary database — a document-oriented NoSQL database accessed via the Motor 3.3 async driver. MongoDB\'s flexible schema enables rapid feature development without rigid migration constraints.', space_before=8, space_after=8)
add_heading('8.1 List of Database Collections', 3)
colls = [
    ('users','All platform users — patients, doctors, admins','500+'),
    ('doctors','Doctor professional profiles and availability','20+'),
    ('appointments','Appointment bookings and status tracking','1000+'),
    ('health_records','Patient medical records (lab, prescription, diagnosis)','2000+'),
    ('prescriptions','Doctor-generated digital prescriptions','500+'),
    ('messages','Real-time chat messages between users','5000+'),
    ('otp_records','Temporary OTP verification records (TTL)','Temporary'),
    ('payments','Razorpay payment transactions','500+'),
    ('insurance_policies','Patient insurance policy records','200+'),
    ('insurance_claims','Insurance claim submissions and status','300+'),
    ('medicine_reminders','Patient medicine schedules','1000+'),
    ('health_goals','Patient health goal tracking','500+'),
    ('vaccination_records','Patient immunization history','300+'),
    ('family_members','Patient family health profiles','400+'),
    ('symptom_diary','Daily symptom logging entries','2000+'),
    ('support_tickets','Admin support ticket management','100+'),
]
add_table(['Collection Name','Purpose','Est. Documents'], colls, [1.6, 3.5, 1.1])
add_para('Table 8.1: MongoDB Collections Overview', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('8.2 Fields and Data Types', 3)
add_para('users Collection:', bold=True, space_before=8, space_after=4)
users_fields = [
    ('_id','ObjectId (PK)','Auto-generated unique identifier'),
    ('full_name','String','User\'s complete name'),
    ('email','String (unique)','Login email — must be unique'),
    ('password_hash','String','bcrypt hashed password (cost 12)'),
    ('role','String (enum)','patient / doctor / admin'),
    ('phone','String','Mobile number (E.164 format)'),
    ('is_active','Boolean','Account active status'),
    ('is_verified','Boolean','Email OTP verification status'),
    ('date_of_birth','Date / null','Patient\'s date of birth'),
    ('gender','String / null','Male / Female / Other'),
    ('address','String / null','Residential address'),
    ('created_at','DateTime','Account creation timestamp (UTC)'),
    ('updated_at','DateTime','Last update timestamp (UTC)'),
]
add_table(['Field','Data Type','Description'], users_fields, [1.5, 1.5, 3.2])
add_para('Table 8.2: users Collection', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)

add_para('doctors Collection:', bold=True, space_after=4)
doc_fields = [
    ('_id','ObjectId (PK)','Unique doctor profile identifier'),
    ('user_id','String (FK → users)','Linked user account ID'),
    ('name','String','Doctor\'s full name with title'),
    ('specialization','String','Medical specialization'),
    ('qualification','String','Degrees (e.g., MBBS, MD)'),
    ('experience_years','Integer','Years of medical experience'),
    ('hospital','String','Current hospital/clinic'),
    ('consultation_fee','Float','Fee per consultation (INR)'),
    ('rating','Float','Average patient rating (0–5)'),
    ('availability','Array[Object]','[{day:"Monday", slots:["09:00","10:00"]}]'),
    ('is_verified','Boolean/null','Admin verification status'),
    ('languages','Array[String]','Languages spoken'),
    ('created_at','DateTime','Profile creation timestamp'),
]
add_table(['Field','Data Type','Description'], doc_fields, [1.5, 1.5, 3.2])
add_para('Table 8.3: doctors Collection', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_para('appointments Collection:', bold=True, space_before=8, space_after=4)
apt_fields = [
    ('_id','ObjectId (PK)','Unique appointment identifier'),
    ('patient_id','String (FK → users)','Patient user ID'),
    ('doctor_id','String (FK → doctors)','Doctor profile ID'),
    ('patient_name','String','Patient name (denormalized)'),
    ('doctor_name','String','Doctor name (denormalized)'),
    ('appointment_date','String (YYYY-MM-DD)','Scheduled date'),
    ('appointment_time','String (HH:MM)','Scheduled time slot'),
    ('appointment_type','String (enum)','video / in-person'),
    ('symptoms','String','Patient-described symptoms'),
    ('status','String (enum)','pending/confirmed/completed/cancelled'),
    ('consultation_fee','Float','Fee charged for this appointment'),
    ('payment_id','String / null','Linked Razorpay payment ID'),
    ('created_at','DateTime','Booking creation timestamp'),
]
add_table(['Field','Data Type','Description'], apt_fields, [1.5, 1.5, 3.2])
add_para('Table 8.4: appointments Collection', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)

add_heading('8.3 Primary and Foreign Keys', 3)
pk_data = [
    ('users','_id (ObjectId)'),('doctors','_id (ObjectId)'),
    ('appointments','_id (ObjectId)'),('health_records','_id (ObjectId)'),
    ('messages','_id (ObjectId)'),('payments','_id (ObjectId)'),
    ('insurance_policies','_id (ObjectId)'),('medicine_reminders','_id (ObjectId)'),
]
add_table(['Collection','Primary Key'], pk_data, [3.0, 3.0])
add_para('Foreign Keys: doctors.user_id → users._id | appointments.patient_id → users._id | appointments.doctor_id → doctors._id | health_records.patient_id → users._id | messages.sender_id → users._id | payments.user_id → users._id | payments.appointment_id → appointments._id | insurance_claims.policy_id → insurance_policies._id', space_before=8, space_after=0)
add_heading('8.4 Relationships', 3)
rels = [
    ('users','doctors','One-to-One'),('users','appointments','One-to-Many'),
    ('doctors','appointments','One-to-Many'),('users','health_records','One-to-Many'),
    ('users','messages','One-to-Many'),('users','payments','One-to-Many'),
    ('users','medicine_reminders','One-to-Many'),('users','health_goals','One-to-Many'),
    ('users','insurance_policies','One-to-Many'),('insurance_policies','insurance_claims','One-to-Many'),
    ('appointments','payments','One-to-One'),
]
add_table(['Parent Collection','Child Collection','Relationship Type'], rels, [2.0, 2.0, 2.0])
add_para('Table 8.7: Collection Relationships', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 9: BACKEND DESCRIPTION
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('9. BACK-END DESCRIPTION', 2)
add_para('The backend of Synora Health is developed using Python 3.11 and the FastAPI framework, organized as a modular REST API with async processing through Motor\'s MongoDB driver, JWT-based authentication middleware, role-based access control, and integration with multiple third-party AI and payment services.', space_before=8, space_after=8)
add_heading('9.1 Database Structure', 3)
add_para('Backend folder structure:', space_after=4)
structure = ('backend/\n├── app/\n│   ├── main.py              ← FastAPI app, CORS, router registration\n'
    '│   ├── config/settings.py   ← Environment variables (Pydantic Settings)\n'
    '│   ├── database/connection.py ← Motor async MongoDB connection\n'
    '│   ├── middleware/auth_middleware.py ← JWT decode, role-based guards\n'
    '│   ├── routes/              ← auth, doctors, appointments, payments, chat, admin\n'
    '│   ├── services/            ← auth_service.py, doctor_service.py\n'
    '│   ├── schemas/             ← Pydantic request/response models\n'
    '│   ├── utils/               ← jwt_utils, password_utils, email_utils, sms_utils\n'
    '│   └── ai/symptom_analyzer.py ← Groq + Gemini AI integration\n'
    '├── requirements.txt\n└── .env                     ← Environment variables')
code_block(structure)
add_heading('9.2 Data Storage and Handling', 3)
add_para('All data operations are performed asynchronously using Motor\'s async MongoDB API. Key utilities: serialize_doc() converts MongoDB documents to JSON-serializable dictionaries; str_to_objectid() safely converts strings to ObjectId; paginate_query() calculates skip/limit for cursor pagination; find_doctor_for_user() searches by user_id first then email fallback.', space_before=8, space_after=8)
add_para('Security in data handling: all passwords bcrypt-hashed before storage; JWT tokens are stateless; OTP codes expire after 10 minutes and deleted after verification; payment verification uses HMAC-SHA256; all queries use parameterized filters to prevent injection.', space_after=0)
add_heading('9.3 Queries and Logic', 3)
add_para('Authentication: Two-step OTP flow, Twilio SMS with email fallback, auto-verification for users with correct passwords. Appointment Booking: availability check → Razorpay order → HMAC verification → appointment creation → WebSocket doctor notification. AI Analysis: structured medical prompt → Groq Llama 3.1 → Gemini fallback → OpenAI fallback. Chat: WebSocket connection pool per room_id, immediate persistence to MongoDB, REST polling for notification badges.', space_before=8, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 10: IMPLEMENTATION
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('10. IMPLEMENTATION', 2)
add_para('Implementation converts the Synora Health system design into a fully functional healthcare platform using Python FastAPI, React 18, MongoDB, and multiple AI/cloud service integrations.', space_before=8, space_after=8)
add_heading('10.1 Module-Wise Explanation', 3)
modules = [
    ('1. Authentication & OTP Module','Handles complete user lifecycle — registration with Twilio SMS OTP, JWT access/refresh token management, auto-verification for users with correct passwords, and role-based redirection. Redux Persist maintains session continuity.'),
    ('2. Doctor Discovery & Booking Module','Patients browse verified doctors, book via calendar date picker and time chip selector, with Razorpay payment integration. Booking confirmation triggers WebSocket notification to doctor.'),
    ('3. AI Health Copilot Module','Role-aware LLM chat powered by Groq Llama 3.1-70B. Doctors receive clinical prompts; patients receive health guidance prompts. Maintains session conversation context.'),
    ('4. AI Symptom Checker Module','Structured medical analysis using multi-step prompting. Results include conditions with confidence, urgency classification, and recommended next steps. Optional image upload triggers Gemini vision analysis.'),
    ('5. Smart Report Analyzer Module','Medical document upload processed through Google Gemini 1.5 Pro vision API. Extracts health parameters, reference ranges, interprets abnormal values. Results stored in health_records with ai_analysis field.'),
    ('6. Real-Time Chat Module','WebSocket-based messaging between patients and doctors. Messages persist to MongoDB immediately. REST polling via GET /chat/rooms detects unread messages. Sidebar badge shows real unread count.'),
    ('7. Video Consultation Module','WebRTC peer-to-peer video calling from patient (/patient/video/:id) and doctor (/doctor/video/:id) routes. Supports local/remote streams, audio/video muting, screen sharing, in-call text chat.'),
    ('8. Payment Gateway Module','Three-step Razorpay integration: GET /payments/key → POST /payments/create-order → POST /payments/verify (HMAC-SHA256). Demo payment modal when keys not configured.'),
    ('9. Medicine Reminder Module','Google Calendar integration generates pre-filled event URLs. Individual CalendarPlus per medicine and bulk "Add All to Calendar". Active/inactive toggle per reminder.'),
    ('10. Health Tracking Suite','Health Goals (5-step wizard, required validation), Vaccination Tracker, Family Health, Symptom Diary, Patient Journey, Insurance management — all backed by real MongoDB data.'),
    ('11. Admin Panel Module','Real-time stats from /admin/stats, clickable stat cards, doctor verification workflow, user management (activate/deactivate/delete), activity logs, support tickets. Notification read state persisted in localStorage.'),
]
for title, desc in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': '); r1.bold = True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(desc); r2.font.name='Times New Roman'; r2.font.size=Pt(12)
add_heading('10.2 Working of the System', 3)
steps = [
    'User Registration & Authentication: OTP via Twilio SMS → OTP verify → JWT tokens → Redux Persist session.',
    'Role-Based Dashboard Access: Protected routes enforce RBAC. Sidebar navigation dynamically rendered by role.',
    'Appointment Booking: Doctor selected → Calendar → Slot → Symptoms → Razorpay → Doctor notified via WebSocket.',
    'AI Health Analysis: Symptoms input → Groq Llama 3.1 → structured response → display with urgency indicators.',
    'Consultation: Doctor confirms → WebRTC video call → prescription written → records updated → appointment completed.',
    'Real-Time Chat: WebSocket delivery → MongoDB persistence → Inbox button shows unread badge → sidebar updates.',
    'Admin Oversight: New doctor registers → admin verifies → verified doctors appear in patient listing.',
    'Health Record Management: Upload → type-specific form → MongoDB save → optional Gemini analysis → stored.',
]
for i, step in enumerate(steps, 1):
    add_para(f'Step {i}: {step}', space_after=4)
add_heading('10.3 Technologies Used', 3)
tech_full = [
    ('Python 3.11','Backend Language','Core server-side logic'),
    ('FastAPI 0.109','API Framework','REST API + WebSocket server'),
    ('Motor 3.3','DB Driver','Async MongoDB operations'),
    ('MongoDB 7.0','Database','Document storage for all data'),
    ('React 18','Frontend Framework','Component-based UI'),
    ('Vite 5.4','Build Tool','Ultra-fast frontend development'),
    ('Tailwind CSS 3.4','CSS Framework','Utility-first responsive styling'),
    ('Redux Toolkit 2.2','State Management','Auth state + appointments'),
    ('Groq SDK (Llama 3.1)','AI Engine','AI Copilot + Symptom Analysis'),
    ('Google Gemini 1.5 Pro','AI Vision','Medical Report Analyzer'),
    ('Twilio','SMS Service','OTP delivery via SMS'),
    ('Razorpay','Payment Gateway','Appointment fee processing'),
    ('WebRTC','Video Protocol','Peer-to-peer video consultation'),
    ('WebSocket','Real-Time','Bidirectional chat communication'),
    ('JWT + bcrypt','Security','Authentication + password hashing'),
]
add_table(['Technology','Category','Purpose in Synora Health'], tech_full, [1.7, 1.5, 3.0])
add_para('Table 10.1: Technologies Used in Synora Health', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 11: SOURCE CODE
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('11. SOURCE CODE', 2)
add_para('Synora Health is implemented across a modular codebase with clear separation of concerns between the Python FastAPI backend and React frontend. Key code snippets from the most important modules are presented below.', space_before=8, space_after=8)
add_heading('11.1 Important Modules and Key Code Snippets', 3)

snippets = [
    ('1. FastAPI Application Entry Point (main.py)',
     'from fastapi import FastAPI\nfrom fastapi.middleware.cors import CORSMiddleware\nfrom contextlib import asynccontextmanager\n\n@asynccontextmanager\nasync def lifespan(app: FastAPI):\n    await connect_db()\n    yield\n    client.close()\n\napp = FastAPI(title="Synora Health API", lifespan=lifespan)\napp.add_middleware(CORSMiddleware,\n    allow_origins=["http://localhost:5173"],\n    allow_credentials=True, allow_methods=["*"], allow_headers=["*"])\n\napp.include_router(auth.router, prefix="/api/auth")\napp.include_router(appointments.router, prefix="/api/appointments")\napp.include_router(payments.router, prefix="/api/payments")\napp.include_router(admin.router, prefix="/api/admin")'),
    ('2. Login Service (auth_service.py)',
     'async def login_user(data: UserLoginRequest) -> dict:\n    user = await db.users.find_one({"email": data.email})\n    if not user:\n        raise HTTPException(401, detail="Invalid credentials")\n    if not verify_password(data.password, user["password_hash"]):\n        raise HTTPException(401, detail="Invalid credentials")\n    if not user.get("is_active"):\n        raise HTTPException(403, detail="Account deactivated")\n    # Auto-verify if password correct\n    if not user.get("is_verified"):\n        await db.users.update_one(\n            {"_id": user["_id"]},\n            {"$set": {"is_verified": True}})\n    serialized = serialize_doc(user)\n    if serialized.get("role") == "doctor":\n        doctor = await db.doctors.find_one({"email": data.email})\n        if doctor:\n            serialized.update({"specialization": doctor.get("specialization"),\n                               "doctor_verified": doctor.get("is_verified", False)})\n    access_token = create_access_token(\n        {"sub": serialized["id"], "role": serialized["role"]})\n    return {"access_token": access_token, "user": serialized}'),
    ('3. AI Symptom Analysis (symptom_analyzer.py)',
     'async def analyze_symptoms(data: dict) -> dict:\n    prompt = f"""You are a medical AI assistant. Analyze:\n    Symptoms: {data["symptoms"]}\n    Duration: {data.get("duration","Unknown")}\n    Severity: {data.get("severity","Moderate")}\n    Provide JSON: possible_conditions, urgency_level,\n    recommended_actions, specialist_required"""\n    try:\n        client = Groq(api_key=settings.GROQ_API_KEY)\n        response = client.chat.completions.create(\n            model="llama-3.1-70b-versatile",\n            messages=[{"role": "user", "content": prompt}],\n            temperature=0.3)\n        return parse_ai_response(response.choices[0].message.content)\n    except Exception:\n        # Fallback to Gemini\n        import google.generativeai as genai\n        genai.configure(api_key=settings.GEMINI_API_KEY)\n        model = genai.GenerativeModel("gemini-1.5-pro")\n        return parse_ai_response(model.generate_content(prompt).text)'),
    ('4. Razorpay Payment Verification (payments.py)',
     'async def verify_payment(data: PaymentVerifyRequest, user=Depends(get_current_user)):\n    expected_sig = hmac.new(\n        settings.RAZORPAY_KEY_SECRET.encode(),\n        f"{data.razorpay_order_id}|{data.razorpay_payment_id}".encode(),\n        hashlib.sha256).hexdigest()\n    if not hmac.compare_digest(expected_sig, data.razorpay_signature):\n        raise HTTPException(400, detail="Payment verification failed")\n    await db.payments.update_one(\n        {"razorpay_order_id": data.razorpay_order_id},\n        {"$set": {"razorpay_payment_id": data.razorpay_payment_id, "status": "verified"}})\n    result = await db.appointments.insert_one({**data.appointment_data,\n        "patient_id": user["id"], "status": "confirmed",\n        "payment_id": data.razorpay_payment_id})\n    return {"appointment_id": str(result.inserted_id), "status": "confirmed"}'),
    ('5. WebSocket Chat Handler (chat.py)',
     'rooms: dict[str, list[WebSocket]] = {}\n\n@router.websocket("/ws/{room_id}")\nasync def websocket_chat(ws: WebSocket, room_id: str):\n    await ws.accept()\n    rooms.setdefault(room_id, []).append(ws)\n    try:\n        while True:\n            data = await ws.receive_json()\n            msg_doc = {"room_id": room_id, "sender_id": data["sender_id"],\n                       "content": data["content"], "timestamp": datetime.utcnow()}\n            await db.messages.insert_one(msg_doc)\n            for conn in rooms.get(room_id, []):\n                if conn != ws:\n                    await conn.send_json(serialize_doc(msg_doc))\n    except WebSocketDisconnect:\n        rooms[room_id].remove(ws)'),
    ('6. Redux Auth Slice with REHYDRATE Fix (authSlice.js)',
     'export const loginUser = createAsyncThunk("auth/login",\n  async (data, { rejectWithValue }) => {\n    try {\n      const res = await api.post("/auth/login", data)\n      toast.success(`Welcome back, ${res.data.user.full_name}!`)\n      return res.data\n    } catch (err) {\n      toast.error(err.response?.data?.detail || "Login failed")\n      return rejectWithValue(err.response?.data)\n    }\n  })\n\nextraReducers: (builder) => {\n  builder\n    // Fix stuck loading spinner from Redux Persist\n    .addCase("persist/REHYDRATE", (state) => { state.loading = false })\n    .addCase(loginUser.pending, (state) => { state.loading = true })\n    .addCase(loginUser.fulfilled, (state, action) => {\n      state.loading = false\n      state.user = action.payload.user\n      state.accessToken = action.payload.access_token })\n    .addCase(loginUser.rejected, (state) => { state.loading = false })\n}'),
    ('7. Dynamic API Base URL (api.js)',
     '// Auto-detects hostname for phone/LAN access\nconst BASE_URL = import.meta.env.VITE_API_URL ||\n  (window.location.hostname === "localhost"\n    ? "http://localhost:8000/api"\n    : `http://${window.location.hostname}:8000/api`)\n\nconst api = axios.create({ baseURL: BASE_URL, timeout: 30000 })\n\n// Attach JWT to every request\napi.interceptors.request.use((config) => {\n  const token = store.getState().auth?.accessToken\n  if (token) config.headers.Authorization = `Bearer ${token}`\n  return config })\n\n// Auto-refresh on 401\napi.interceptors.response.use(response => response,\n  async (error) => {\n    if (error.response?.status === 401 && !error.config._retry) {\n      error.config._retry = true\n      const refresh = store.getState().auth?.refreshToken\n      if (refresh) {\n        const res = await axios.post(`${BASE_URL}/auth/refresh`,\n          { refresh_token: refresh })\n        store.dispatch(setTokens(res.data))\n        error.config.headers.Authorization = `Bearer ${res.data.access_token}`\n        return api(error.config) }\n    }\n    return Promise.reject(error) })'),
]
for title, code in snippets:
    add_para(title, bold=True, space_before=12, space_after=4)
    code_block(code)
    add_para(f'Figure 11.1.{snippets.index((title,code))+1}: {title.split("(")[0].strip()}', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 12: TESTING
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('12. TESTING AND IMPLEMENTATION', 2)
add_para('Testing ensures Synora Health functions correctly, securely, and reliably across all modules and user roles. The system was tested using Unit Testing, Integration Testing, System Testing, and User Acceptance Testing.', space_before=8, space_after=8)
add_heading('12.1 Unit Testing', 3)
ut = [
    ('UT-01','Register new patient with valid data','OTP sent, user pending verification','Pass ✓'),
    ('UT-02','Verify correct 6-digit OTP','Account created, JWT tokens returned','Pass ✓'),
    ('UT-03','Login with valid credentials','Access + refresh tokens issued','Pass ✓'),
    ('UT-04','Login with wrong password','HTTP 401 Invalid credentials','Pass ✓'),
    ('UT-05','Login with deactivated account','HTTP 403 Account deactivated','Pass ✓'),
    ('UT-06','Book appointment with valid slot','Appointment created with pending status','Pass ✓'),
    ('UT-07','Doctor approves appointment','Status updated to confirmed','Pass ✓'),
    ('UT-08','Update doctor availability (min 1 day)','Availability saved to MongoDB','Pass ✓'),
    ('UT-09','Update with zero working days','HTTP 400 validation error','Pass ✓'),
    ('UT-10','Verify valid Razorpay HMAC signature','Payment verified, appointment booked','Pass ✓'),
    ('UT-11','Verify tampered HMAC signature','HTTP 400 verification failed','Pass ✓'),
    ('UT-12','Submit symptoms to Groq API','Structured analysis response returned','Pass ✓'),
    ('UT-13','Add Lab Report record','Record saved with correct type fields','Pass ✓'),
    ('UT-14','Fetch platform statistics (admin)','Real counts from MongoDB returned','Pass ✓'),
    ('UT-15','Admin verifies doctor profile','is_verified set to true in doctors collection','Pass ✓'),
    ('UT-16','Send message via WebSocket','Message persisted and broadcast to room','Pass ✓'),
    ('UT-17','Register with already verified email','HTTP 400 Email already registered','Pass ✓'),
    ('UT-18','Doctor accesses insurance endpoint','HTTP 200 (not 403) — role allowed','Pass ✓'),
]
add_table(['Test ID','Test Description','Expected Result','Status'], ut, [0.6, 2.2, 2.3, 0.9])
add_para('Table 12.1: Unit Testing Results', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('12.2 Integration Testing', 3)
it = [
    ('IT-01','Auth + MongoDB','Complete OTP registration flow','User stored in MongoDB after OTP verify','Pass ✓'),
    ('IT-02','Auth + Twilio','SMS OTP delivery to real phone','6-digit OTP received via SMS','Pass ✓'),
    ('IT-03','Appointments + Payment','Book appointment with Razorpay','Payment verified → appointment created','Pass ✓'),
    ('IT-04','Appointments + WebSocket','Appointment booked → doctor notified','Doctor dashboard shows new appointment','Pass ✓'),
    ('IT-05','Symptom Checker + Groq','Submit symptoms → AI analysis','Conditions, urgency, actions returned','Pass ✓'),
    ('IT-06','AI + Gemini Fallback','Groq unavailable → Gemini fallback','Analysis returned from Gemini API','Pass ✓'),
    ('IT-07','Chat + MongoDB','WebSocket message → DB persistence','Message visible in chat history on reload','Pass ✓'),
    ('IT-08','Doctors + Admin','Doctor registers → admin verifies','Doctor appears in patient listing after verify','Pass ✓'),
    ('IT-09','Health Records + Gemini','Upload medical image → AI analysis','Parameters extracted and stored with record','Pass ✓'),
    ('IT-10','Medicine + Google Cal.','Click Add to Calendar button','Google Calendar opens with pre-filled event','Pass ✓'),
    ('IT-11','Redux + API Interceptor','Access token expired → auto-refresh','New token fetched, request retried silently','Pass ✓'),
    ('IT-12','Admin + Users + MongoDB','Admin fetches 100 users','All users returned without pagination cap','Pass ✓'),
]
add_table(['Test ID','Modules Integrated','Test Scenario','Expected Result','Status'], it, [0.6, 1.4, 1.6, 1.7, 0.6])
add_para('Table 12.2: Integration Testing Scenarios', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
add_heading('12.3 System Testing', 3)
st = [
    ('ST-01','Patient full registration and login flow','Account created, dashboard accessible','As expected','Pass ✓'),
    ('ST-02','Doctor registration → admin verification → patient listing','Doctor visible to patients after verify','As expected','Pass ✓'),
    ('ST-03','Complete appointment booking with payment','Appointment confirmed, both dashboards updated','As expected','Pass ✓'),
    ('ST-04','AI Symptom Analysis end-to-end','Results displayed within 5 seconds','3.2s average','Pass ✓'),
    ('ST-05','Video consultation initiation (WebRTC)','Video call established','Established on LAN','Pass ✓'),
    ('ST-06','Real-time chat patient ↔ doctor','Messages delivered instantly','As expected','Pass ✓'),
    ('ST-07','Admin doctor verification workflow','Doctor verified/rejected correctly','As expected','Pass ✓'),
    ('ST-08','Medicine reminder + Google Calendar link','Calendar opens with medicine event','As expected','Pass ✓'),
    ('ST-09','Smart report upload + Gemini analysis','Parameters extracted from medical image','As expected','Pass ✓'),
    ('ST-10','Platform access from mobile phone','App fully functional on phone browser','Working (IP:5173)','Pass ✓'),
    ('ST-11','Admin stat cards click → tab navigation','Clicking Total Patients opens patients tab','As expected','Pass ✓'),
    ('ST-12','Session persistence after page refresh','User stays logged in after F5','As expected','Pass ✓'),
]
add_table(['Test ID','Test Description','Expected Output','Actual Output','Status'], st, [0.6, 1.8, 1.5, 1.3, 0.8])
add_para('Table 12.3: System Testing Scenarios', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

chapter_header('Synora Health')
add_heading('12.4 User Acceptance Testing (UAT)', 3)
uat = [
    ('Registration and OTP verification was straightforward','Patient','Positive ✓'),
    ('Doctor listing and availability clearly displayed','Patient','Positive ✓'),
    ('Appointment booking with payment was smooth','Patient','Positive ✓'),
    ('AI Symptom Checker provided helpful results','Patient','Positive ✓'),
    ('AI Health Copilot answered medical questions accurately','Patient','Positive ✓'),
    ('Medicine reminder calendar integration was useful','Patient','Positive ✓'),
    ('Health goals tracking with progress was motivating','Patient','Positive ✓'),
    ('Doctor dashboard provided complete patient overview','Doctor','Positive ✓'),
    ('Appointment approval and rejection was simple','Doctor','Positive ✓'),
    ('Prescription writing module was professional','Doctor','Positive ✓'),
    ('Patient insurance visible inside patient modal','Doctor','Positive ✓'),
    ('Task & Reminders with priority selection was useful','Doctor','Positive ✓'),
    ('Admin stats cards with real data are informative','Admin','Positive ✓'),
    ('Doctor verification workflow was clear and easy','Admin','Positive ✓'),
    ('Overall interface was professional and user-friendly','All Roles','Positive ✓'),
]
add_table(['Observation / Feature Tested','User Role','Result'], uat, [3.5, 1.2, 1.1])
add_para('Table 12.4: User Acceptance Testing Observations', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 13: RESULTS
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('13. RESULTS AND DISCUSSION', 2)
add_para('Synora Health was successfully developed and tested as a fully functional AI-powered healthcare platform. All 13 major modules were implemented, tested, and verified to work correctly with real data from the MongoDB database.', space_before=8, space_after=8)
add_heading('13.1 System Outputs', 3)
out1 = [
    ('OTP delivery channel','SMS (Twilio) → Email fallback → Demo mode'),
    ('OTP expiry','10 minutes from generation'),
    ('JWT access token validity','60 minutes'),
    ('JWT refresh token validity','7 days'),
    ('Password security','bcrypt hash, cost factor 12'),
]
add_table(['Parameter','Result'], out1, [3.0, 3.0])
add_para('Table 13.1: Authentication & Registration Output', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
add_para('AI Symptom Analysis Sample — Input: "Severe headache, blurred vision, nausea for 3 days"', bold=True, space_after=4)
out2 = [
    ('Possible Condition 1','Migraine (75% probability)'),
    ('Possible Condition 2','Hypertensive Episode (60% probability)'),
    ('Possible Condition 3','Tension Headache (45% probability)'),
    ('Urgency Level','Urgent – Consult doctor within 24 hours'),
    ('Specialist Required','Neurologist / General Physician'),
    ('Response Time','2.8 seconds (Groq ultra-fast inference)'),
]
add_table(['Output Parameter','Sample Result'], out2, [2.5, 3.5])
add_para('Table 13.2: AI Symptom Analysis Sample Output', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
add_heading('13.2 Observations', 3)
obs = [
    ('AI Response Quality','Groq Llama 3.1-70B provided clinically relevant responses with clear urgency classification. Delivered in under 3.5 seconds on average.'),
    ('Real-Time Performance','WebSocket messages delivered in under 100ms on LAN. Chat Inbox polling (30-second) effectively detected unread messages without page refresh.'),
    ('Payment Gateway','Razorpay HMAC-SHA256 verification successfully prevented payment tampering. Demo modal maintained development workflow.'),
    ('RBAC','JWT-based role control effectively prevented unauthorized access across all protected routes.'),
    ('Database Performance','Motor async driver maintained sub-100ms query response times. Email-based fallback lookup successfully linked seeded doctors.'),
    ('Mobile Access','Platform accessed and tested on Android Chrome via WiFi after Windows Firewall configuration. Dynamic API URL correctly routed to computer IP.'),
    ('Redux Persist Fix','REHYDRATE action handler prevents stuck loading spinner when previous session ended mid-request with loading:true persisted.'),
]
for title, desc in obs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': '); r1.bold=True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(desc); r2.font.name='Times New Roman'; r2.font.size=Pt(12)
add_heading('13.3 Performance Discussion', 3)
perf = [
    ('API response time (standard)','< 500ms','~120ms avg','Exceeded ✓'),
    ('AI symptom analysis','< 5 sec','2.8–4.2s','Met ✓'),
    ('WebSocket message delivery','< 200ms','< 100ms (LAN)','Exceeded ✓'),
    ('Dashboard load time','< 2 sec','~1.1s','Exceeded ✓'),
    ('MongoDB query time','< 200ms','~80ms avg','Exceeded ✓'),
    ('Mobile accessibility','Full feature','Full feature on LAN','Met ✓'),
]
add_table(['Performance Metric','Target','Achieved','Result'], perf, [2.0, 1.2, 1.3, 1.4])
add_para('Table 13.3: Performance Discussion Summary', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
page_break()

# ════════════════════════════════════════════════════════════
#  CH 14-18
# ════════════════════════════════════════════════════════════
chapter_header('Synora Health')
add_heading('14. LIMITATIONS', 2)
add_heading('14.1 Current Constraints of the System', 3)
lims = [
    ('No HTTPS / SSL in Development','Mobile browsers block geolocation and camera features on HTTP. Nearby Places and camera-based report scanning are limited on mobile browsers without HTTPS deployment.'),
    ('AI Responses Not Medically Certified','AI Health Copilot and Symptom Checker responses are for informational purposes only — not a substitute for professional medical diagnosis. Mandatory disclaimer displayed with all AI outputs.'),
    ('Local-Only Deployment','Currently running on localhost. Cloud deployment (Railway, Vercel, MongoDB Atlas) not completed for this version, limiting public internet accessibility.'),
    ('Video Consultation on LAN Only','WebRTC peer-to-peer video calling works reliably on local networks. STUN/TURN server configuration for internet-based NAT traversal not yet implemented.'),
    ('No Real-Time AI Cost Monitoring','API call costs to Groq, Gemini, and OpenAI are not tracked or limited per user, which could cause unexpected costs in production.'),
    ('No Push Notifications','Notification delivery relies on 30-second polling rather than true push notifications, causing up to 30-second notification delivery delay.'),
    ('Payment in Test Mode','Razorpay configured in test mode. Live payment processing requires additional KYC verification and Razorpay business account.'),
]
for title, desc in lims:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': '); r1.bold=True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(desc); r2.font.name='Times New Roman'; r2.font.size=Pt(12)
page_break()

chapter_header('Synora Health')
add_heading('15. FUTURE ENHANCEMENTS', 2)
add_heading('15.1 Possible Improvements & Additional Features', 3)
future = [
    ('Cloud Deployment with HTTPS','Deploy backend to Railway or AWS EC2, frontend to Vercel, database to MongoDB Atlas. Configure SSL/TLS for HTTPS, enabling full mobile feature access including geolocation and camera.'),
    ('STUN/TURN Server for Global Video Calls','Integrate Coturn STUN/TURN server to enable WebRTC video consultations across the internet, removing the LAN-only limitation.'),
    ('Firebase Push Notifications','Replace polling with Firebase Cloud Messaging (FCM) for real-time push notifications on web and mobile, including background notifications.'),
    ('Mobile Application (React Native)','Develop iOS and Android native apps using React Native, sharing existing business logic and API layer with native push notifications and offline viewing.'),
    ('ABDM / FHIR Integration','Integrate with Ayushman Bharat Digital Mission for standardized health record exchange using FHIR (HL7) format for Indian healthcare interoperability.'),
    ('Advanced AI Diagnostic Models','Train specialized medical ML models on Indian patient datasets for accurate disease prediction, drug interaction checking, and personalized treatment recommendations.'),
    ('Multi-Language Support','Add Hindi, Gujarati, Tamil, and other regional language support to make the platform accessible to India\'s diverse population.'),
    ('Health Wearable Integration','Connect with Apple Health, Google Fit, and Fitbit APIs to automatically import health metrics (heart rate, steps, sleep) into the patient dashboard.'),
    ('AI-Powered Smart Scheduling','ML-based scheduling that considers doctor availability, patient history, urgency level to optimize appointment booking for all parties.'),
    ('Telemedicine Prescription Validation','Integrate with government e-prescription validation for legally valid digital prescriptions verifiable by pharmacies.'),
]
for i, (title, desc) in enumerate(future, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{i}. {title}: '); r1.bold=True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(desc); r2.font.name='Times New Roman'; r2.font.size=Pt(12)
page_break()

chapter_header('Synora Health')
add_heading('16. CONCLUSION', 2)
add_heading('16.1 Summary of the Project', 3)
add_para('"Synora Health" was successfully designed, developed, and tested as a comprehensive, AI-powered healthcare management platform during a 6-month internship at Coding House, Surat. The platform integrates FastAPI, React 18, MongoDB, Groq Llama 3.1, Google Gemini, WebRTC, WebSocket, Razorpay, and Twilio into a unified digital healthcare ecosystem serving patients, doctors, and administrators.', space_before=8, space_after=8)
add_para('The platform successfully delivers: a complete Patient Portal with 15+ features including AI Symptom Analysis, AI Health Copilot, Smart Report Analyzer, appointment booking with payment, video consultation, real-time chat, health records, medicine reminders, health goals, insurance, vaccination tracking, emergency SOS, and nearby hospital finder; a comprehensive Doctor Dashboard; a powerful Admin Panel with real data; three AI engines with intelligent fallback chains; and secure OTP authentication with JWT session management.', space_after=8)
add_para('All testing phases — Unit, Integration, System, and UAT — passed successfully. Performance metrics exceeded targets with API response times averaging 120ms, AI analysis delivered in under 4 seconds, and WebSocket messages delivered in under 100ms.', space_after=0)
add_heading('16.2 Learning Outcomes', 3)
learn = [
    ('1. Advanced Full-Stack Development','Deep practical experience in Python FastAPI with async/await patterns and React 18 with Redux Toolkit, Vite, and Tailwind CSS. Understanding the complete request lifecycle from frontend to database.'),
    ('2. Real-World AI Integration','Integrating multiple LLM providers (Groq, Gemini, OpenAI) with fallback chains. Prompt engineering for medical contexts and handling production-grade AI applications.'),
    ('3. Real-Time Web Technologies','WebSocket-based bidirectional chat and WebRTC peer-to-peer video consultation — critical modern skills rarely covered in academic curriculum.'),
    ('4. Payment Gateway Security','Razorpay HMAC-SHA256 verification, order lifecycle management, and payment fraud prevention through server-side verification.'),
    ('5. Security Engineering','bcrypt password hashing, JWT token rotation, OTP expiry management, CORS configuration, and role-based access control in practice.'),
    ('6. MongoDB & NoSQL Design','Flexible document-oriented schema design for healthcare data with 16+ collections, proper reference relationships, and the Motor async driver.'),
    ('7. Software Engineering','Modular architecture, environment-based configuration, error handling, fallback strategies, version control, and git merge conflict resolution.'),
    ('8. Healthcare Domain Knowledge','Medical workflows, patient data privacy, FHIR data standards, telemedicine regulations, and ethical considerations of AI in medical contexts.'),
]
for title, desc in learn:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': '); r1.bold=True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(desc); r2.font.name='Times New Roman'; r2.font.size=Pt(12)
page_break()

chapter_header('Synora Health')
add_heading('17. BIBLIOGRAPHY / REFERENCES', 2)
add_heading('17.1 Books, Websites, and Other Sources', 3)
add_para('Books:', bold=True, space_before=8, space_after=6)
books = [
    '[1] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. MIT Press, 2016.',
    '[2] E. Alpaydin, Introduction to Machine Learning, 4th ed. MIT Press, 2020.',
    '[3] A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 3rd ed. O\'Reilly Media, 2022.',
    '[4] E. Balagurusamy, Programming in Python, 2nd ed. McGraw Hill Education India, 2019.',
    '[5] R. Natarajan, Machine Learning and Deep Learning Using Python. Wiley India Pvt. Ltd., 2021.',
]
for b in books: add_para(b, space_after=6)
add_para('Official Documentation and Websites:', bold=True, space_before=8, space_after=6)
refs = [
    '[6] FastAPI Documentation. https://fastapi.tiangolo.com/. Accessed: May 2026.',
    '[7] MongoDB Documentation. https://www.mongodb.com/docs/. Accessed: May 2026.',
    '[8] React Documentation. https://react.dev/. Accessed: May 2026.',
    '[9] Groq API Documentation. https://console.groq.com/docs/. Accessed: May 2026.',
    '[10] Google Generative AI Documentation. https://ai.google.dev/docs. Accessed: May 2026.',
    '[11] Razorpay Payment Gateway Documentation. https://razorpay.com/docs/. Accessed: May 2026.',
    '[12] Twilio SMS Documentation. https://www.twilio.com/docs/. Accessed: May 2026.',
    '[13] Redux Toolkit Documentation. https://redux-toolkit.js.org/. Accessed: May 2026.',
    '[14] Tailwind CSS Documentation. https://tailwindcss.com/docs/. Accessed: May 2026.',
    '[15] Pydantic v2 Documentation. https://docs.pydantic.dev/. Accessed: May 2026.',
    '[16] WebRTC Documentation. https://developer.mozilla.org/en-US/docs/Web/API/WebRTC_API. Accessed: May 2026.',
    '[17] Python Software Foundation. https://docs.python.org/3/. Accessed: May 2026.',
    '[18] Microsoft Visual Studio Code. https://code.visualstudio.com/. Accessed: May 2026.',
    '[19] GitHub Inc. https://github.com/. Accessed: May 2026.',
    '[20] OpenStreetMap Nominatim API. https://nominatim.openstreetmap.org/. Accessed: May 2026.',
]
for r in refs: add_para(r, space_after=6)
page_break()

chapter_header('Synora Health')
add_heading('18. APPENDIX', 2)
add_heading('18.1 Additional Data, Extended Code, or Supporting Material', 3)
add_para('Environment Configuration (.env):', bold=True, space_before=8, space_after=4)
env_code = ('# MongoDB\nMONGODB_URL=mongodb://localhost:27017\nDATABASE_NAME=aihealthcare\n\n'
    '# JWT\nSECRET_KEY=your-256-bit-secret-key\nACCESS_TOKEN_EXPIRE_MINUTES=60\n\n'
    '# AI APIs\nGROQ_API_KEY=gsk_xxxxxxxxxxxxxxxxxxxxx\nGEMINI_API_KEY=AIzaSyxxxxxxxxxxxxxxxxx\n\n'
    '# Twilio SMS\nTWILIO_ACCOUNT_SID=ACxxxxxxxxxxxxxxxxxx\nTWILIO_AUTH_TOKEN=xxxxxxxx\nTWILIO_FROM_NUMBER=+1xxxxxxxxxx\n\n'
    '# Razorpay\nRAZORPAY_KEY_ID=rzp_test_xxxxxxxxxx\nRAZORPAY_KEY_SECRET=xxxxxxxx\n\n'
    '# CORS\nCORS_ORIGINS=["http://localhost:5173"]')
code_block(env_code)
add_para('Server Start Commands:', bold=True, space_before=8, space_after=4)
start_code = ('# Backend (FastAPI)\ncd backend\nvenv\\Scripts\\uvicorn app.main:app --reload --host 0.0.0.0 --port 8000\n\n'
    '# Frontend (React + Vite)\ncd frontend\nnpm run dev\n\n'
    '# Access Points\nFrontend:  http://localhost:5173\nAPI Docs:  http://localhost:8000/docs\nMobile:    http://192.168.0.109:5173')
code_block(start_code)
add_para('Project Folder Structure:', bold=True, space_before=8, space_after=4)
folder_data = [
    ('backend/app/main.py','FastAPI application entry point'),
    ('backend/app/routes/','All API route handlers (auth, doctors, appointments, payments, chat, admin)'),
    ('backend/app/services/','Business logic services (auth_service, doctor_service)'),
    ('backend/app/schemas/','Pydantic request/response models'),
    ('backend/app/utils/','JWT, password, email, SMS, helper utilities'),
    ('backend/app/ai/','AI integration (Groq, Gemini, symptom analyzer)'),
    ('backend/app/database/','MongoDB Motor connection management'),
    ('frontend/src/pages/','All page components (30+ pages)'),
    ('frontend/src/components/','Reusable UI components (Navbar, Sidebar, Chat, AIChatbot)'),
    ('frontend/src/redux/','Redux store, slices (auth, appointments, doctors, ai)'),
    ('frontend/src/services/api.js','Axios instance with JWT interceptor and refresh logic'),
    ('frontend/src/routes/','Protected route component with RBAC'),
    ('frontend/src/App.jsx','React Router configuration, all route definitions'),
]
add_table(['File / Folder','Purpose'], folder_data, [2.5, 3.7])
add_para('Plagiarism / Originality Report:', bold=True, space_before=12, space_after=4)
add_para('[Figure 18.1.1 – Plagiarism Check Report (Paperpal + Turnitin)]\n\nSimilarity range: 3% – 9%\n\nSources: Research Papers on Healthcare AI Systems (2%), FastAPI / React Official Documentation references (1%), Medical terminology and standard clinical terms (<1%)\n\n→ Replace this with your actual Paperpal/Turnitin plagiarism report screenshot', italic=True, space_after=0)

# ════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════
output_path = r'c:\Users\Dell\AI\SYNORA_HEALTH_REPORT.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
size_kb = __import__('os').path.getsize(output_path) // 1024
print(f'File size: {size_kb} KB')
